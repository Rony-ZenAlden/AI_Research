"""Text extraction from uploaded files.

Each extractor is dependency-light and streaming-friendly so big uploads don't
blow our 4GB VPS memory budget. CSV is capped at 5,000 rows to be safe.
"""
from __future__ import annotations

import csv
import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# Extensions we accept on the upload endpoint — keep in sync with serializer.
SUPPORTED_EXTENSIONS = frozenset({"pdf", "docx", "txt", "md", "markdown", "csv"})

CSV_ROW_CAP = 5000


class UnsupportedFileType(ValueError):
    pass


class ExtractionError(RuntimeError):
    pass


def extract_text(file_path: str | Path) -> tuple[str, dict[str, Any]]:
    """Dispatch to the right extractor based on file extension.

    Returns ``(text, metadata)`` — text is plain UTF-8; metadata is extractor-
    specific (page counts, paragraph counts, etc.) and merged into the
    Document's metadata JSON.
    """
    path = Path(file_path)
    ext = path.suffix.lower().lstrip(".")
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(f"unsupported extension: .{ext}")

    if ext == "pdf":
        return _extract_pdf(path)
    if ext == "docx":
        return _extract_docx(path)
    if ext == "csv":
        return _extract_csv(path)
    # txt, md, markdown
    return _extract_plain(path)


def _extract_pdf(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages_text: list[str] = []
        for page in reader.pages:
            text = (page.extract_text() or "").strip()
            if text:
                pages_text.append(text)
        full = "\n\n".join(pages_text)
        meta = {
            "extractor": "pypdf",
            "page_count": len(reader.pages),
            "extracted_pages": len(pages_text),
        }
        return full, meta
    except Exception as e:
        raise ExtractionError(f"PDF extraction failed: {e}") from e


def _extract_docx(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from docx import Document as DocxDocument

        d = DocxDocument(str(path))
        paragraphs = [p.text for p in d.paragraphs if p.text and p.text.strip()]
        # Pull table cells too — a lot of docx content lives in tables.
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        paragraphs.append(cell.text.strip())
        full = "\n\n".join(paragraphs)
        meta = {"extractor": "python-docx", "paragraph_count": len(paragraphs)}
        return full, meta
    except Exception as e:
        raise ExtractionError(f"DOCX extraction failed: {e}") from e


def _extract_plain(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return text, {"extractor": "plain", "encoding": "utf-8"}
    except Exception as e:
        raise ExtractionError(f"Text read failed: {e}") from e


def _extract_csv(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        rows_out: list[str] = []
        headers: list[str] = []
        truncated = False
        count = 0
        with path.open("r", encoding="utf-8", errors="replace", newline="") as f:
            reader = csv.reader(f)
            headers = next(reader, []) or []
            if headers:
                rows_out.append(" | ".join(str(h) for h in headers))
            for row in reader:
                rows_out.append(" | ".join(str(c) for c in row))
                count += 1
                if count >= CSV_ROW_CAP:
                    rows_out.append(f"[... truncated after {CSV_ROW_CAP} rows ...]")
                    truncated = True
                    break
        text = "\n".join(rows_out)
        meta = {
            "extractor": "csv",
            "row_count": count,
            "columns": headers,
            "truncated": truncated,
        }
        return text, meta
    except Exception as e:
        raise ExtractionError(f"CSV read failed: {e}") from e
